#!/usr/bin/env python3
"""TAKE FIVE 사용 설명서 — 워드 파일."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Emu, Pt, RGBColor, Twips
from docx.oxml.ns import nsdecls
from lxml import etree

INK = RGBColor(0x18, 0x17, 0x14)
MUTED = RGBColor(0x6B, 0x66, 0x5C)
CREAM = RGBColor(0xF4, 0xF1, 0xEA)
REC = RGBColor(0xC4, 0x5C, 0x4A)
STEEL = RGBColor(0x5A, 0x62, 0x6C)
CARD = RGBColor(0x13, 0x13, 0x16)
RULE = "B8B3A8"
SOFT = "E8E4DA"

OUT = Path("/workspace/TAKE-FIVE-사용설명서.docx")
ART = Path("/workspace/artifacts/TAKE-FIVE-사용설명서.docx")


def set_run_font(run, name="Malgun Gothic", size=11, bold=False, color=INK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def set_cell_borders(cell, color="D6D1C6", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = tcMar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tcMar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")


def set_table_width(table, width_cm):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    twips = int(width_cm * 567)
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")


def para_space(p, before=0, after=8, line=1.22):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TAKE FIVE  ·  ")
    set_run_font(run, size=9, color=MUTED)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    set_run_font(r2, size=9, color=MUTED)
    r2._r.append(fldChar1)
    r2._r.append(instr)
    r2._r.append(fldChar2)
    run3 = p.add_run("  /  ")
    set_run_font(run3, size=9, color=MUTED)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r4 = p.add_run()
    set_run_font(r4, size=9, color=MUTED)
    r4._r.append(fldChar3)
    r4._r.append(instr2)
    r4._r.append(fldChar4)


def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def add_bottom_border(p, color=RULE, sz="12"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def write_cell(cell, text, *, size=10.5, bold=False, color=INK, align="left", fill=None):
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p, 2, 2, 1.15)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_kv_table(doc, rows, col_widths=(3.6, 12.8)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 16.4)
    for i, (k, v) in enumerate(rows):
        fill = "F7F4EC" if i % 2 == 0 else "FFFFFF"
        c0, c1 = table.rows[i].cells
        write_cell(c0, k, size=10.5, bold=True, fill=fill)
        write_cell(c1, v, size=10.5, fill=fill)
        c0.width = Cm(col_widths[0])
        c1.width = Cm(col_widths[1])
    doc.add_paragraph()
    return table


def heading(doc, text, n):
    p = doc.add_paragraph()
    para_space(p, 16, 8, 1.15)
    add_bottom_border(p, RULE, "8")
    r1 = p.add_run(f"{n}  ")
    set_run_font(r1, name="Georgia", size=11, bold=True, color=REC)
    r2 = p.add_run(text)
    set_run_font(r2, name="Malgun Gothic", size=16, bold=True, color=INK)
    return p


def subhead(doc, text):
    p = doc.add_paragraph()
    para_space(p, 12, 4, 1.15)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=INK)
    return p


def body(doc, text, after=8):
    p = doc.add_paragraph()
    para_space(p, 0, after, 1.32)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def bullet(doc, text, hang=True):
    p = doc.add_paragraph()
    para_space(p, 1, 3, 1.28)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r0 = p.add_run("·  ")
    set_run_font(r0, size=11, color=REC)
    r1 = p.add_run(text)
    set_run_font(r1, size=11, color=INK)
    return p


def callout(doc, title, text, fill="F4EDE6"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 16.4)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, 140, 140, 180, 180)
    set_cell_borders(cell, "D9CFC3", "8")
    p0 = cell.paragraphs[0]
    para_space(p0, 0, 4, 1.2)
    r = p0.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=REC)
    p1 = cell.add_paragraph()
    para_space(p1, 0, 0, 1.28)
    r2 = p1.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def numbered(doc, n, title, text):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, 16.4)
    no_table_borders(table)
    c0, c1 = table.rows[0].cells
    c0.width = Cm(1.6)
    c1.width = Cm(14.8)
    shade_cell(c0, "181714")
    set_cell_margins(c0, 80, 80, 60, 60)
    set_cell_margins(c1, 60, 80, 140, 40)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p0, 2, 2, 1)
    r = p0.add_run(n)
    set_run_font(r, name="Georgia", size=12, bold=True, color=CREAM)
    p1 = c1.paragraphs[0]
    para_space(p1, 0, 2, 1.15)
    r1 = p1.add_run(title)
    set_run_font(r1, size=12, bold=True, color=INK)
    p2 = c1.add_paragraph()
    para_space(p2, 0, 0, 1.25)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=MUTED)
    spacer = doc.add_paragraph()
    para_space(spacer, 0, 6, 1)


def faq_item(doc, q, a):
    p = doc.add_paragraph()
    para_space(p, 10, 2, 1.2)
    r = p.add_run("Q.  " + q)
    set_run_font(r, size=11, bold=True, color=INK)
    p2 = doc.add_paragraph()
    para_space(p2, 0, 4, 1.28)
    r2 = p2.add_run(a)
    set_run_font(r2, size=11, color=INK)


def check_item(doc, n, text):
    p = doc.add_paragraph()
    para_space(p, 2, 4, 1.25)
    r0 = p.add_run("☐   ")
    set_run_font(r0, size=12, color=REC)
    r1 = p.add_run(f"{n}.  ")
    set_run_font(r1, name="Georgia", size=11, bold=True, color=MUTED)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    add_page_number(section)

    # ---- Cover ----
    for _ in range(3):
        sp = doc.add_paragraph()
        para_space(sp, 0, 0, 1)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(kicker, 0, 6, 1)
    rk = kicker.add_run("HACKATHON DEMO STUDIO")
    set_run_font(rk, name="Georgia", size=10, color=REC)
    rk.font.all_caps = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(title, 4, 4, 1)
    rt = title.add_run("TAKE FIVE")
    set_run_font(rt, name="Georgia", size=42, bold=True, color=INK)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(rule, 2, 10, 1)
    add_bottom_border(rule, "C45C4A", "18")
    rr = rule.add_run(" ")
    set_run_font(rr, size=4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(sub, 6, 18, 1.3)
    rs = sub.add_run("해커톤 시연 스튜디오 사용 설명서")
    set_run_font(rs, size=16, color=INK)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(tag, 0, 22, 1.4)
    rtg = tag.add_run(
        "화면을 먼저 찍고, 말은 나중에 입힙니다.\n"
        "5분 이내 시연 영상을 대본 · 나레이션 · 자막까지 한곳에서 완성하세요."
    )
    set_run_font(rtg, size=11, color=MUTED)

    add_kv_table(
        doc,
        [
            ("대상", "해커톤 팀, 데모데이 발표자"),
            ("결과물", "MP4 시연 영상, 자막 SRT, 대본 TXT"),
            ("길이", "5분 이내 (그 전에 끝나도 됩니다)"),
            ("언어", "한국어"),
            ("버전", "TAKE FIVE  ·  2026"),
        ],
    )

    callout(
        doc,
        "한 줄 요약",
        "찍기 → 대본 → 목소리 → 내보내기. 미리보기 창에서 화면 녹화가 막히면 "
        "이미 찍은 영상을 올리거나 샘플로 먼저 흐름을 타 보세요.",
    )

    # ---- 1 ----
    heading(doc, "TAKE FIVE가 하는 일", "01")
    body(
        doc,
        "TAKE FIVE는 해커톤 시연 영상을 급하게 만들어야 할 때 쓰는 스튜디오입니다. "
        "화면을 먼저 찍거나 올리고, 화면에 맞춰 대본을 쓴 뒤, AI 여자·남자 목소리 또는 "
        "내 목소리로 나레이션을 입히고, 자막을 적용할지 정한 다음 MP4로 저장합니다.",
    )
    body(
        doc,
        "작품 제목은 화면 위쪽 입력칸에서 바꿀 수 있습니다. 기본 제목은 ‘해커톤 시연’입니다. "
        "같은 브라우저에서는 제목, 대본, 자막, 목소리 선택이 기억됩니다.",
    )

    subhead(doc, "전체 흐름")
    numbered(
        doc,
        "01",
        "찍기",
        "화면 · 카메라 · 업로드 · 샘플 중 하나로 클립을 넣습니다. 얼굴 캠은 선택할 수 있습니다.",
    )
    numbered(
        doc,
        "02",
        "대본",
        "타임코드에 맞춰 할 말을 적거나 txt · srt · vtt 파일을 올립니다. 자막도 여기서 만들 수 있습니다.",
    )
    numbered(
        doc,
        "03",
        "목소리",
        "AI 여자/남자 목소리로 생성하거나, 영상을 보며 직접 읽습니다.",
    )
    numbered(
        doc,
        "04",
        "내보내기",
        "볼륨 · 자막 적용/미적용 · 자르기를 정한 뒤 MP4로 저장합니다.",
    )

    # ---- 2 ----
    heading(doc, "시작하기 전에", "02")
    bullet(doc, "컴퓨터는 Chrome, Edge, Safari를 권장합니다.")
    bullet(doc, "화면 녹화 · 카메라 · 마이크를 쓰려면 브라우저 권한을 허용하세요.")
    bullet(
        doc,
        "미리보기 창(채팅 옆 작은 화면)에서는 화면 · 카메라 녹화가 막히는 경우가 많습니다.",
    )
    bullet(
        doc,
        "그때는 ‘영상 올리기’로 이미 찍은 파일을 넣거나, ‘샘플’로 먼저 연습하세요.",
    )
    bullet(
        doc,
        "아이폰에서 영상을 올릴 때는 파일 종류가 비어 있어도 .mp4 / .mov 이면 받습니다.",
    )
    bullet(doc, "5분은 최대 길이입니다. 1~2분이어도 됩니다. 준비되면 바로 종료하세요.")

    callout(
        doc,
        "미리보기가 목업처럼 보일 때",
        "미리보기 창은 실제 앱입니다. 다만 바깥 화면을 캡처하는 권한은 이 창에서 자주 거절됩니다. "
        "‘영상 올리기’ 카드 전체를 누르면 파일 선택이 열립니다. 숨은 ‘열기’ 버튼을 찾을 필요는 없습니다.",
    )

    # ---- 3 ----
    heading(doc, "찍기", "03")
    body(
        doc,
        "앱 위쪽 01 · 찍기 단계입니다. 화면을 먼저 찍습니다. 말은 나중에 입힙니다. 5분 이내면 됩니다. "
        "4분 30초가 되면 타이머가 빨간색으로 바뀌고, 5분이 되면 자동으로 멈춥니다.",
    )

    subhead(doc, "클립을 넣는 네 가지 방법")
    add_kv_table(
        doc,
        [
            (
                "화면",
                "탭 또는 창을 녹화합니다. ‘얼굴 같이 찍기’가 켜져 있으면 화면 한쪽에 얼굴이 겹칩니다.",
            ),
            ("카메라만", "얼굴 풀샷만 녹화합니다. 발표자 클로즈업이 필요할 때 씁니다."),
            (
                "영상 올리기",
                "이미 찍은 파일을 넣습니다. 미리보기에서 녹화가 안 될 때 가장 확실합니다. mp4, mov, webm을 받습니다.",
            ),
            (
                "샘플",
                "12초짜리 연습 클립입니다. 대본 · 목소리 · 자막 흐름을 바로 타 볼 수 있습니다. 얼굴 캠이 붙는 구성도 미리 보입니다.",
            ),
        ],
        col_widths=(3.4, 13.0),
    )

    subhead(doc, "얼굴 같이 찍기")
    body(
        doc,
        "유튜브 라이브처럼 화면 한쪽에 촬영자 얼굴을 겹칩니다. 스위치로 끄고 켤 수 있습니다. 기본은 켜져 있습니다.",
    )
    bullet(doc, "위치: 왼위, 오른위, 왼아래, 오른아래")
    bullet(doc, "모양: 둥글게 / 네모")
    bullet(doc, "크기: 작게 / 보통")
    bullet(
        doc,
        "마이크 같이 녹음: 화면 소리와 내 목소리를 함께 담습니다. 나중에 원본 소리를 끌 수도 있습니다.",
    )
    body(
        doc,
        "얼굴을 끄면 화면만 녹화됩니다. 얼굴만 크게 찍으려면 ‘카메라만’을 쓰세요. "
        "이 미리보기 창에서는 실제 카메라가 안 붙을 수 있으니, 샘플로 레이아웃만 먼저 확인하면 됩니다.",
    )

    subhead(doc, "녹화 중")
    bullet(doc, "카운트다운 3 · 2 · 1 뒤에 시작됩니다.")
    bullet(doc, "언제든 ‘녹화 종료’를 누르면 끝납니다. 5분이 되면 자동으로 멈춥니다.")
    bullet(doc, "5분보다 긴 클립은 내보낼 때 잘립니다.")
    bullet(
        doc,
        "클립을 버리고 다시 시작하려면 ‘클립 버리고 다시’를 누르세요. 대본과 목소리도 함께 지워집니다.",
    )

    # ---- 4 ----
    heading(doc, "대본", "04")
    body(
        doc,
        "앱 위쪽 02 · 대본 단계입니다. 화면에 맞춰 말을 적습니다. 줄을 누르면 그 시각으로 영상이 이동합니다. "
        "한 줄은 1~3문장이 적당합니다. 구어체로 쓰세요.",
    )
    bullet(doc, "현재 시각에 큐: 지금 보고 있는 장면에 빈 줄을 넣습니다.")
    bullet(
        doc,
        "5분 뼈대: 오프닝 – 문제 – 솔루션 – 라이브 데모 – 차별점 – 클로징 틀을 넣습니다. 문장만 바꾸면 됩니다.",
    )
    bullet(
        doc,
        "대본 올리기: .txt, .srt, .vtt, .md 파일을 받습니다. [0:12] 한 줄 처럼 시각이 있으면 그 시각에 맞춥니다. 없으면 문단을 나눠 길이에 맞춥니다.",
    )
    bullet(
        doc,
        "AI로 쓰기: 작품 이름, 문제, 솔루션, 데모에서 보여줄 것, 말투(담백 / 진지 / 경쾌)를 적으면 한국어 큐를 받습니다.",
    )
    bullet(
        doc,
        "자막 자동 생성: 소리 또는 대본을 자막 줄로 나눕니다. 나레이션 대본은 지우지 않습니다.",
    )

    body(
        doc,
        "위쪽에 예상 발화 시간과 영상 길이가 보입니다. 말이 영상보다 길면 ‘조금 줄여 주세요’가 뜹니다. "
        "줄이 서로 겹치면 ‘말이 겹칩니다’와 ‘겹침 풀기’가 나옵니다. 겹침 풀기를 누르면 간격이 맞춰집니다.",
    )

    callout(
        doc,
        "잘 되는 대본",
        "한 줄은 1~3문장, 구어체로. ‘지금부터 시연하겠습니다’ 같은 사회자 멘트는 빼 주세요. "
        "데모 구간에 시간을 가장 많이 주세요. 화면이 바뀌는 순간에 줄을 나누면 나레이션이 장면과 맞습니다.",
    )

    # ---- 5 ----
    heading(doc, "목소리", "05")
    body(
        doc,
        "앱 위쪽 03 · 목소리 단계입니다. AI 목소리로 생성하거나, 영상을 보며 직접 읽습니다. 여자 / 남자를 고르면 목록이 바뀝니다. "
        "‘들어보기’로 샘플을 듣고, ‘전체에 입히기’로 대본 전체에 입힙니다.",
    )

    subhead(doc, "AI 목소리")
    add_kv_table(
        doc,
        [
            (
                "여자",
                "이브(기본, 따뜻하고 또렷한), 아라(맑고 밝은), 루나(차분한), 아이리스(부드러운), "
                "오로라(경쾌한), 리오라(신뢰감), 셀레스트(우아한), 카리나(또렷하고 분명한), 우르사(무게감 있는)",
            ),
            (
                "남자",
                "오리온(기본, 시네마틱), 레오(무게감 있는), 렉스(또렷하고 분명한), 살(부드럽고 균형 잡힌), "
                "페르세우스(신뢰감), 아틀라스(자신감 있는), 럭스(차분한), 리겔(프로페셔널), 자간(드라마틱), 헬릭스(다이나믹)",
            ),
            ("속도", "0.70× ~ 1.50×. 기본은 1.00×입니다. 말이 빠르면 0.90×로 내려 보세요."),
        ],
        col_widths=(2.8, 13.6),
    )
    body(
        doc,
        "생성할 때 줄이 겹치지 않게 간격을 다시 맞춥니다. 이미 만들어 둔 자막은 지우지 않습니다. "
        "미리듣기를 연속으로 눌러도 목소리가 겹치지 않습니다.",
    )
    callout(
        doc,
        "AI를 쓸 수 없을 때",
        "이 환경에서 AI 목소리를 쓸 수 없으면 안내가 뜹니다. 그때는 ‘내가 읽기’로 진행하세요.",
    )

    subhead(doc, "내가 읽기")
    bullet(doc, "지금 읽을 줄이 크게 보입니다. 영상이 그 시각으로 맞춰집니다.")
    bullet(doc, "리허설: 녹화하지 않고 화면만 재생합니다. 텔레프롬프터처럼 읽으면 됩니다.")
    bullet(
        doc,
        "영상과 함께 녹음: 영상이 처음부터 재생되며 마이크가 켜집니다. 끝나면 ‘녹음 끝내기’를 누르세요.",
    )
    bullet(doc, "새 클립을 넣으면 예전 나레이션은 지워집니다. 대본은 길이에 맞춰 다시 배치됩니다.")

    # ---- 6 ----
    heading(doc, "자막", "06")
    body(
        doc,
        "자막은 대본과 따로 둡니다. 나레이션 대본을 지우지 않고, 영상에 올릴 글만 다룹니다. "
        "소리를 인식하거나, 소리가 없으면 대본을 짧은 한국어 줄로 나눕니다.",
    )
    bullet(
        doc,
        "자막 자동 생성: 대본 단계, 목소리 단계, 내보내기 단계에서 누를 수 있습니다. 목소리를 입힌 뒤에는 나레이션을 우선합니다.",
    )
    bullet(
        doc,
        "적용: 미리보기 아래쪽에 자막이 뜨고, 저장 영상에도 새겨집니다. 자막 줄이 없으면 대본을 자막으로 씁니다.",
    )
    bullet(
        doc,
        "미적용: 미리보기와 저장 영상에서 자막이 빠집니다. SRT 파일은 그대로 받을 수 있습니다.",
    )
    bullet(doc, "줄마다 글자를 고치거나 지울 수 있습니다.")
    bullet(doc, "지우기: 자막 줄만 지웁니다. 대본은 남습니다.")
    bullet(
        doc,
        "자막 자동 생성을 누르면 적용으로 바뀝니다. 영상에 넣기 싫으면 내보내기에서 다시 미적용을 고르세요.",
    )

    # ---- 7 ----
    heading(doc, "내보내기", "07")
    body(
        doc,
        "앱 위쪽 04 · 내보내기 단계입니다. 목소리를 섞고, 자막을 넣을지 정한 뒤 MP4로 저장합니다. "
        "나레이션이 없어도 영상만 잘라 저장할 수 있습니다.",
    )
    add_kv_table(
        doc,
        [
            (
                "원본 소리 끄기",
                "녹화할 때 담긴 화면 소리를 끕니다. 끄면 나레이션만 들립니다. AI 목소리를 입혔다면 켜 두는 편이 깔끔합니다.",
            ),
            ("원본 볼륨", "원본 소리를 켤 때만 나옵니다."),
            ("나레이션 볼륨", "AI 또는 내 목소리 크기입니다."),
            (
                "자막",
                "적용이면 미리보기와 저장 영상에 자막이 들어갑니다. 미적용이면 빠집니다.",
            ),
            ("자르기", "앞뒤를 잘라 내보낼 구간을 정합니다. 침묵이 길면 앞뒤를 조금 자르세요."),
        ],
        col_widths=(3.6, 12.8),
    )

    subhead(doc, "저장 파일")
    bullet(
        doc,
        "MP4로 저장: 심사 페이지에 바로 올리는 형식입니다. 이 브라우저가 MP4를 못 만들면 WebM으로 내려갑니다.",
    )
    bullet(
        doc,
        "자막 SRT: 영상에 새기지 않고 자막 파일만 받을 때 씁니다. 자막을 미적용해도 받을 수 있습니다.",
    )
    bullet(doc, "대본 TXT: 타임코드가 붙은 대본입니다. 발표 리허설용으로 쓰기 좋습니다.")

    callout(
        doc,
        "저장이 끝나기까지",
        "영상 길이만큼 섞은 뒤 MP4로 바꿉니다. 5분 클립은 몇 분이 걸릴 수 있습니다. "
        "진행 막대가 보이면 기다리세요. 끝나면 파일이 내려갑니다.",
    )

    # ---- 8 ----
    heading(doc, "폰에서 쓸 때", "08")
    bullet(doc, "화면이 세로로 쌓입니다. 위가 영상, 아래가 단계입니다.")
    bullet(
        doc,
        "영상 올리기 · 대본 올리기는 카드/버튼 전체를 누르세요. 숨은 ‘열기’가 아니라 실제 파일 선택입니다.",
    )
    bullet(
        doc,
        "화면 · 카메라 녹화는 미리보기와 아이폰에서 자주 막힙니다. 폰 카메라로 찍은 뒤 영상 올리기가 가장 잘 됩니다.",
    )
    bullet(
        doc,
        "글자만 보이고 화면이 비면 잠시 기다리거나, 가로로 돌려 보세요. 샘플을 넣으면 영상이 채워집니다.",
    )
    bullet(
        doc,
        "얼굴 같이 찍기는 화면 녹화가 될 때 붙습니다. 폰 미리보기에서는 레이아웃만 고르고, 실제 얼굴은 컴퓨터 브라우저에서 찍는 편이 낫습니다.",
    )

    subhead(doc, "폰에서 권장하는 순서")
    numbered(doc, "1", "영상 올리기", "갤러리에 있는 시연 영상을 고릅니다. 없으면 샘플로 흐름만 먼저 탑니다.")
    numbered(doc, "2", "대본", "줄을 직접 적거나 대본 올리기로 txt/srt를 넣습니다. 겹치면 겹침 풀기.")
    numbered(doc, "3", "자막", "자막 자동 생성을 누른 뒤, 넣을지 말지 적용/미적용으로 정합니다.")
    numbered(doc, "4", "목소리", "여자 또는 남자 AI 목소리를 입히거나, 화면을 보며 직접 읽습니다.")
    numbered(doc, "5", "저장", "원본 소리 끄기, 자막 여부를 확인한 뒤 MP4로 저장합니다.")

    # ---- 9 ----
    heading(doc, "자주 묻는 질문", "09")
    faq_item(
        doc,
        "5분이 되기 전에 끝나도 되나요?",
        "됩니다. 준비되면 바로 녹화 종료를 누르세요. 5분은 최대 길이입니다.",
    )
    faq_item(
        doc,
        "미리보기에서 화면이 안 찍혀요. 목업인가요?",
        "실제 앱입니다. 미리보기 창이 화면 캡처를 막는 경우가 많습니다. 영상 올리기 또는 샘플을 쓰세요. "
        "따로 연 브라우저 탭이나 배포된 앱에서는 녹화가 됩니다.",
    )
    faq_item(
        doc,
        "폰에서 ‘열기’가 안 돼요.",
        "카드 전체를 누르세요. 영상 올리기 · 대본 올리기는 숨은 작은 버튼이 아니라 카드 자체가 파일 선택입니다.",
    )
    faq_item(
        doc,
        "말이 겹쳐 들립니다.",
        "대본 단계에서 ‘겹침 풀기’를 누른 뒤 목소리를 다시 입히세요. AI 생성은 줄이 겹치지 않게 붙입니다. "
        "샘플은 짧은 클립이라 대본이 길면 겹칠 수 있습니다.",
    )
    faq_item(
        doc,
        "여자 목소리는 어디 있나요?",
        "목소리 단계 → AI 목소리 → 여자. 기본은 이브입니다. 들어보기로 먼저 들어 보세요.",
    )
    faq_item(
        doc,
        "자막을 영상에는 안 넣고 싶어요.",
        "내보내기에서 미적용을 고르세요. 미리보기에서도 빠집니다. SRT는 따로 받을 수 있습니다.",
    )
    faq_item(
        doc,
        "영상을 바꿨더니 예전 목소리가 남아요?",
        "새 클립을 넣으면 예전 나레이션은 지워집니다. 대본은 길이에 맞춰 다시 배치됩니다.",
    )
    faq_item(
        doc,
        "MP4 대신 WebM이 내려왔습니다.",
        "그 브라우저가 H.264/AAC를 못 만들 때입니다. Chrome이나 Edge에서 다시 저장하거나, 받은 WebM을 변환하세요.",
    )
    faq_item(
        doc,
        "작업이 사라졌나요?",
        "같은 브라우저에서는 제목, 대본, 자막, 목소리 선택이 기억됩니다. 영상 · 나레이션 파일도 이 기기에 남습니다. "
        "다른 폰 · 컴퓨터와는 공유되지 않습니다.",
    )
    faq_item(
        doc,
        "권한 안내가 떠요.",
        "‘권한을 허용해 주세요. 미리보기 창에서는 화면 녹화가 막힐 수 있습니다.’가 보이면 "
        "영상을 올리거나 샘플로 이어가면 됩니다.",
    )

    # ---- 10 ----
    heading(doc, "해커톤 당일 체크리스트", "10")
    check_item(doc, "1", "화면을 찍었거나, 영상을 올렸거나, 샘플로 흐름을 확인했다.")
    check_item(doc, "2", "대본이 화면 타임코드와 맞는다. 말이 겹치면 겹침 풀기를 했다.")
    check_item(doc, "3", "AI 여자/남자 목소리 또는 내 목소리를 입혔다. 들어 보고 속도 · 볼륨을 만졌다.")
    check_item(doc, "4", "자막을 넣을지 정했다. 적용이면 미리보기에서 한 번 확인했다.")
    check_item(doc, "5", "원본 소리를 끌지 정했고, 앞뒤 공백을 잘랐다.")
    check_item(doc, "6", "MP4로 저장했다. 재생해 보고 소리 · 자막을 확인했다.")
    check_item(doc, "7", "심사 페이지에 올렸다. 필요하면 SRT도 첨부했다.")

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(end, 28, 4, 1.3)
    add_bottom_border(end, "C45C4A", "12")
    re = end.add_run("화면을 먼저 찍고, 말은 나중에.  5분이면 충분합니다.")
    set_run_font(re, size=12, italic=True, color=MUTED)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(foot, 8, 0, 1)
    rf = foot.add_run("TAKE FIVE  ·  해커톤 시연 스튜디오")
    set_run_font(rf, name="Georgia", size=10, color=MUTED)

    props = doc.core_properties
    props.title = "TAKE FIVE 사용 설명서"
    props.author = "TAKE FIVE"
    props.subject = "해커톤 시연 스튜디오 사용 설명서"
    props.language = "ko-KR"

    OUT.write_bytes(b"")  # ensure parent
    doc.save(str(OUT))
    ART.parent.mkdir(parents=True, exist_ok=True)
    ART.write_bytes(OUT.read_bytes())
    print("wrote", OUT, OUT.stat().st_size)
    print("copied", ART, ART.stat().st_size)


if __name__ == "__main__":
    build()
